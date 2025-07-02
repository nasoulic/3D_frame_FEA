import os
import glob
import re
import math
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
# from docx2pdf import convert


class EngineeringReportGenerator:
    def __init__(self, data_dir, output_name="Engineering_Report"):
        self.data_dir = data_dir
        self.output_name = output_name
        self.properties_files = {}
        self.stresses_files = {}
        self.nodal_disp_files = {}
        self.images = {}

    def collect_files(self):
        for file in glob.glob(os.path.join(self.data_dir, "*.dat")):
            fname = os.path.basename(file)
            if "_properties.dat" in fname:
                name = fname.split("_properties")[0]
                self.properties_files[name] = file
            elif "stresses_output_" in fname:
                loadcase = fname.split("stresses_output_")[1].split(".dat")[0]
                self.stresses_files[loadcase] = file
            elif "nodal_displacement_" in fname:
                loadcase = fname.split("nodal_displacement_")[1].split(".dat")[0]
                self.nodal_disp_files[loadcase] = file

        for file in glob.glob(os.path.join(self.data_dir, "*.png")):
            fname = os.path.basename(file)
            if "_stress_" in fname:
                loadcase = fname.split("_stress_")[-1].replace(".png", "")
            else:
                parts = fname.split("_")
                loadcase = parts[-1].replace(".png", "")
            if loadcase not in self.images:
                self.images[loadcase] = []
            self.images[loadcase].append(file)

    def parse_properties_file(self, filepath):
        props = {}
        with open(filepath) as f:
            for line in f:
                if "=" in line:
                    key, val = line.split("=")
                    key = key.strip()
                    val = val.split("#")[0].strip()
                    try:
                        val = float(val)
                        val = round(val, 1)
                    except:
                        if val.lower() == "none":
                            val = None
                    props[key] = val
        return props

    def parse_stresses_file(self, filepath):
        elements = []
        with open(filepath) as f:
            block = {}
            for line in f:
                line = line.strip()
                if line.startswith("Stresses at element"):
                    if block:
                        elements.append(block)
                        block = {}
                    block["ID"] = int(re.findall(r'\d+', line)[0])
                elif "Element group" in line:
                    block["group"] = line.split(":")[1].strip()
                elif ":" in line:
                    k, v = line.split(":")
                    block[k.strip()] = float(v.strip().split(" ")[0])
            if block:
                elements.append(block)
        return elements

    def parse_nodal_displacement_file(self, filepath):
        nodes = []
        with open(filepath) as f:
            for line in f:
                if line.startswith("Node"):
                    parts = line.split(":")
                    node_id = int(parts[0].split()[1])
                    disp = re.findall(r'[-+]?\d*\.\d+e[+-]\d+', line)
                    Ux, Uy, Uz, Rx, Ry, Rz = map(float, disp)
                    nodes.append({
                        "ID": node_id,
                        "Ux": Ux,
                        "Uy": Uy,
                        "Uz": Uz,
                        "Rx": Rx,
                        "Ry": Ry,
                        "Rz": Rz,
                        "TotalDisp": math.sqrt(Ux**2 + Uy**2 + Uz**2)
                    })
        return nodes

    def add_image_with_caption(self, doc, image_path):
        doc.add_picture(image_path, width=Inches(5))
        image_name = os.path.basename(image_path).replace(".png", "")
        doc.add_paragraph(f"Figure: {image_name}", style='Caption')

    def add_properties_table(self, doc, filepath, name):
        doc.add_heading(f"Properties: {name}", level=2)

        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].paragraphs[0].add_run("Property").bold = True
        hdr[1].paragraphs[0].add_run("Value").bold = True
        hdr[2].paragraphs[0].add_run("Description / Units").bold = True

        with open(filepath) as f:
            for line in f:
                if "=" in line:
                    key, rest = line.split("=")
                    key = key.strip()
                    val_comment = rest.split("#")
                    val = val_comment[0].strip()
                    comment = val_comment[1].strip() if len(val_comment) > 1 else ""

                    try:
                        val = float(val)
                        val = f"{val:.1f}"
                    except:
                        if val.lower() == "none":
                            val = "None"

                    row = table.add_row().cells
                    row[0].text = key
                    row[1].text = str(val)
                    row[2].text = comment

        self.set_table_borders(table)

    def add_stress_table(self, doc, elements):
        doc.add_heading("Stresses Table - Full", level=3)
        table = doc.add_table(rows=1, cols=8)
        hdr = table.rows[0].cells
        headers = ["Element ID", "Group", "σ_axial [MPa]", "σ_bend_y [MPa]",
                   "σ_bend_z [MPa]", "τ_torsion [MPa]", "σ_total [MPa]", "σ_vonMises [MPa]"]
        for i, h in enumerate(headers):
            hdr[i].paragraphs[0].add_run(h).bold = True

        for el in elements:
            row = table.add_row().cells
            row[0].text = str(el["ID"])
            row[1].text = el["group"]
            row[2].text = f'{el["sigma_axial"]:.2f}'
            row[3].text = f'{el["sigma_bending_y"]:.2f}'
            row[4].text = f'{el["sigma_bending_z"]:.2f}'
            row[5].text = f'{el["tau_torsion"]:.2f}'
            row[6].text = f'{el["sigma_total"]:.2f}'

            vm = math.sqrt(el["sigma_axial"] ** 2 + el["sigma_bending_y"] ** 2 +
                           el["sigma_bending_z"] ** 2 + 3 * el["tau_torsion"] ** 2)
            row[7].text = f'{vm:.2f}'
            el["vonMises"] = vm

        self.set_table_borders(table)

    def add_max_stress_summary(self, doc, elements):
        max_el = max(elements, key=lambda x: x["vonMises"])
        doc.add_heading("Maximum von Mises Stress", level=3)
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        headers = ["Element ID", "Group", "Max σ_vonMises  [MPa]"]
        for i, h in enumerate(headers):
            hdr[i].paragraphs[0].add_run(h).bold = True
        row = table.add_row().cells
        row[0].text = str(max_el["ID"])
        row[1].text = max_el["group"]
        row[2].text = f'{max_el["vonMises"]:.2f}'
        self.set_table_borders(table)

    def add_nodal_displacement_table(self, doc, nodes):
        doc.add_heading("Total Nodal Displacement - Full", level=3)
        table = doc.add_table(rows=1, cols=8)
        hdr = table.rows[0].cells
        headers = ["Node ID", "Ux [mm]", "Uy [mm]", "Uz [mm]", "Rx [rad]", "Ry [rad]", "Rz [rad]", "Total Disp [mm]"]
        for i, h in enumerate(headers):
            hdr[i].paragraphs[0].add_run(h).bold = True

        for node in nodes:
            row = table.add_row().cells
            row[0].text = str(node["ID"])
            row[1].text = f'{node["Ux"]:.3f}'
            row[2].text = f'{node["Uy"]:.3f}'
            row[3].text = f'{node["Uz"]:.3f}'
            row[4].text = f'{node["Rx"]:.3f}'
            row[5].text = f'{node["Ry"]:.3f}'
            row[6].text = f'{node["Rz"]:.3f}'
            row[7].text = f'{node["TotalDisp"]:.3f}'

        self.set_table_borders(table)

    def add_max_displacement_summary(self, doc, nodes):
        max_node = max(nodes, key=lambda x: x["TotalDisp"])
        doc.add_heading("Maximum Total Displacement", level=3)
        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].paragraphs[0].add_run("Node ID").bold = True
        hdr[1].paragraphs[0].add_run("Max Total Disp [mm]").bold = True
        row = table.add_row().cells
        row[0].text = str(max_node["ID"])
        row[1].text = f'{max_node["TotalDisp"]:.3f}'
        self.set_table_borders(table)

    def set_table_borders(self, table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)
        tblPr.append(tblBorders)

    def compute_von_mises(self, elements):
        for el in elements:
            vm = math.sqrt(el["sigma_axial"] ** 2 + el["sigma_bending_y"] ** 2 +
                        el["sigma_bending_z"] ** 2 + 3 * el["tau_torsion"] ** 2)
            el["vonMises"] = vm


    def generate(self):
        self.collect_files()
        doc = Document()
        doc.add_heading("Engineering Report", 0)

        # Properties sections
        for name, file in self.properties_files.items():
            self.add_properties_table(doc, file, name)

        all_loadcases = set(self.stresses_files.keys()).union(
            self.nodal_disp_files.keys()).union(self.images.keys())

        for loadcase in sorted(all_loadcases):
            doc.add_heading(f'Loadcase: {loadcase}', level=1)

            if loadcase in self.images:
                doc.add_heading('Images', level=2)
                for img in self.images[loadcase]:
                    self.add_image_with_caption(doc, img)

            if loadcase in self.stresses_files:
                doc.add_heading('Stress Results', level=2)
                elements = self.parse_stresses_file(self.stresses_files[loadcase])
                self.compute_von_mises(elements)
                self.add_max_stress_summary(doc, elements)
                self.add_stress_table(doc, elements)
                
                
            if loadcase in self.nodal_disp_files:
                doc.add_heading('Nodal Displacements', level=2)
                nodes = self.parse_nodal_displacement_file(self.nodal_disp_files[loadcase])
                self.add_max_displacement_summary(doc, nodes)
                self.add_nodal_displacement_table(doc, nodes)

        docx_path = f"{self.output_name}.docx"
        doc.save(docx_path)
        # convert(docx_path, f"{self.output_name}.pdf")
        # print(f"Report saved: {docx_path} & {self.output_name}.pdf")

if __name__ == "__main__":
    gen = EngineeringReportGenerator(os.getcwd())
    gen.generate()
